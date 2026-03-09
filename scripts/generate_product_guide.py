#!/usr/bin/env python3
"""
Generate Patent Space MCP Product Guide as a professional Word document.
Output: /tmp/Patent_Space_MCP_Product_Guide.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Constants ──────────────────────────────────────────────────────
OUTPUT_PATH = "/tmp/Patent_Space_MCP_Product_Guide.docx"
DARK_BLUE = RGBColor(0x1B, 0x36, 0x5D)
HEADER_BG = "D6E4F0"
ALT_ROW_BG = "F2F7FB"
WHITE_BG = "FFFFFF"
LIGHT_GRAY_BG = "F0F0F0"
FONT_BODY = "Calibri"
FONT_CODE = "Consolas"


# ── Helper Functions ───────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with val, sz, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "1B365D")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def styled_paragraph(doc, text, style_name=None, font_name=FONT_BODY,
                     font_size=Pt(11), bold=False, color=None,
                     alignment=None, space_after=Pt(6), space_before=Pt(0)):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    if style_name:
        p.style = doc.styles[style_name]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    return p


def add_heading_styled(doc, text, level=1):
    """Add a heading with custom styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_BODY
        run.font.color.rgb = DARK_BLUE
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_BODY)
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return h


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with header shading and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_BODY)
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), FONT_BODY)
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)

    # Apply column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    return table


def add_code_block(doc, text):
    """Simulate a code block using a single-cell table with gray background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    for line in text.split("\n"):
        p = cell.paragraphs[0] if not cell.paragraphs[0].text and cell.paragraphs[0].runs == [] else cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = FONT_CODE
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
    set_cell_shading(cell, LIGHT_GRAY_BG)
    doc.add_paragraph()  # spacing


def add_metric_box(doc, metrics):
    """Add a key metrics highlight box as a table."""
    table = doc.add_table(rows=len(metrics), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(metrics):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        run0 = p0.add_run(value)
        run0.font.name = FONT_BODY
        run0.font.size = Pt(14)
        run0.font.bold = True
        run0.font.color.rgb = DARK_BLUE
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p1 = c1.paragraphs[0]
        run1 = p1.add_run(label)
        run1.font.name = FONT_BODY
        run1.font.size = Pt(11)

        set_cell_shading(c0, HEADER_BG)
        set_cell_shading(c1, "EEF3FA")

        c0.width = Cm(5)
        c1.width = Cm(11)
    doc.add_paragraph()


def add_page_break(doc):
    """Insert a page break."""
    doc.add_page_break()


def add_body_text(doc, text):
    """Add body text paragraph."""
    return styled_paragraph(doc, text, font_size=Pt(11), space_after=Pt(8))


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_BODY)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p


# ── Section Builders ───────────────────────────────────────────────

def build_title_page(doc):
    """Page 1: Title page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    styled_paragraph(doc, "Patent Space MCP", font_size=Pt(36), bold=True,
                     color=DARK_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     space_after=Pt(12))

    # Horizontal rule
    p_hr = doc.add_paragraph()
    p_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hr = p_hr.add_run("_" * 60)
    run_hr.font.color.rgb = DARK_BLUE
    run_hr.font.size = Pt(11)

    styled_paragraph(doc, "プロダクト紹介資料", font_size=Pt(28), bold=True,
                     color=DARK_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     space_after=Pt(24))

    styled_paragraph(
        doc,
        "世界初・67ツール搭載 AI特許分析プラットフォーム",
        font_size=Pt(16), color=RGBColor(0x4A, 0x6F, 0xA5),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48)
    )

    styled_paragraph(doc, "2026年3月", font_size=Pt(14),
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    styled_paragraph(doc, "v1.0", font_size=Pt(12),
                     color=RGBColor(0x80, 0x80, 0x80),
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    add_page_break(doc)


def build_executive_summary(doc):
    """Pages 2-3: Executive Summary."""
    add_heading_styled(doc, "Executive Summary", level=1)

    add_body_text(
        doc,
        "Patent Space MCPは、150,000,000件超のグローバル特許データベース"
        "を基盤に、67の専門分析ツールを"
        "Model Context Protocol (MCP)経由で提供する世界初のAI特許分析"
        "プラットフォームです。PatSnap（年間300万円）やDerwent Innovation"
        "（年間200万円）と同等以上の分析能力を、オープンソースで実現します。"
    )

    doc.add_paragraph()

    add_heading_styled(doc, "Key Metrics", level=2)

    add_metric_box(doc, [
        ("グローバル特許データベース", "150,000,000件超"),
        ("専門分析ツール", "67ツール"),
        ("技術クラスタ", "607クラスタ"),
        ("事前計算Startabilityスコア", "10,267,405件"),
        ("企業名寄せエントリ", "20,736件"),
        ("事前計算分析の応答時間", "サブ秒"),
    ])

    add_heading_styled(doc, "プラットフォームの特徴", level=2)

    features = [
        "MCPプロトコル対応: Claude Desktop、Cursor等のAIツールからネイティブ接続",
        "事前計算アーキテクチャ: 10M+のStartabilityスコアを事前計算し、サブ秒レスポンスを実現",
        "67の専門ツール: 基本検索からゲーム理論分析、Black-Scholesリアルオプション評価まで網羅",
        "グローバル対応: 150M件超の全世界特許をカバー（JP 26M + US/EP/CN/KR等）",
        "オープンソース: 商用DB（年間200-300万円）と同等機能を無償で提供",
        "拡張可能: カスタム技術カテゴリの定義、AI分類、監視アラートなど",
    ]
    for f in features:
        add_bullet(doc, f)

    add_page_break(doc)


def build_data_foundation(doc):
    """Pages 4-5: Data Foundation."""
    add_heading_styled(doc, "データ基盤", level=1)

    add_body_text(
        doc,
        "Patent Space MCPのデータ基盤は、Google BigQueryから取得した特許メタデータを"
        "SQLiteデータベースに格納し、高速な分析クエリを実現するアーキテクチャです。"
    )

    add_heading_styled(doc, "データパイプライン", level=2)
    add_body_text(
        doc,
        "Google Patents Public Data (BigQuery) からParquet形式でエクスポートし、"
        "Hetznerサーバー上の650GB+ SQLiteデータベースに格納しています。"
    )

    add_heading_styled(doc, "データベース統計", level=2)

    add_styled_table(doc,
        ["テーブル", "レコード数", "説明"],
        [
            ["patents", "150,088,362", "グローバル特許メタデータ（JP/US/EP/CN/KR等）"],
            ["patent_cpc", "44,800,000", "特許-CPC分類コード対応"],
            ["patent_assignees", "30,400,000", "特許-出願人対応"],
            ["firm_tech_vectors", "27,800", "企業技術ベクトル（64次元）"],
            ["startability_surface", "10,267,405", "事前計算Startabilityスコア"],
            ["tech_clusters", "607", "技術クラスタ定義"],
            ["patent_cluster_mapping", "76,266", "特許-クラスタマッピング"],
            ["patent_research_data", "171,500,000", "研究用拡張データ"],
            ["entity_resolution", "20,736", "企業名寄せマッピング"],
        ]
    )

    doc.add_paragraph()

    add_heading_styled(doc, "追加データソース", level=2)

    add_styled_table(doc,
        ["データソース", "件数", "説明"],
        [
            ["SEP宣言 (ETSI ISLD)", "950+", "標準必須特許の宣言データ"],
            ["PTAB審判", "7,605", "米国特許審判データ"],
            ["特許訴訟", "74,629", "特許訴訟ケースデータ"],
            ["GDELTメディアシグナル", "46企業", "主要企業のメディア露出・トーン分析"],
        ]
    )

    doc.add_paragraph()

    add_heading_styled(doc, "グローバル拡張（取込中）", level=2)
    add_body_text(
        doc,
        "現在、グローバル特許133M件（14,216 Parquetファイル、約2.2TB）の取込を進めています。"
        "取込完了後は、米国(USPTO)、欧州(EPO)、世界知的所有権機関(WIPO)の特許もカバーし、"
        "クロスボーダー分析が本格的に可能になります。"
    )

    add_page_break(doc)


def build_tool_list(doc):
    """Pages 6-9: 67 Tool Feature List."""
    add_heading_styled(doc, "67ツール一覧", level=1)

    add_body_text(
        doc,
        "Patent Space MCPが提供する67の専門分析ツールを、カテゴリ別に紹介します。"
        "各ツールはMCPプロトコル経由でAIアシスタントから直接呼び出し可能です。"
    )

    categories = [
        ("基本検索・分析 (Basic Search & Analysis)", [
            ("patent_search", "キーワード、CPC、出願人による特許検索"),
            ("patent_detail", "特許の詳細情報取得"),
            ("firm_patent_portfolio", "企業の特許ポートフォリオ分析"),
            ("entity_resolve", "企業名の正規化・名寄せ"),
            ("tech_landscape", "技術領域の出願動向・トップ出願人"),
            ("applicant_network", "共同出願ネットワーク分析"),
            ("patent_compare", "複数企業の特許ポートフォリオ比較"),
        ]),
        ("Startability分析 (Technology Entry Analysis)", [
            ("tech_fit", "企業×技術クラスタの適合度計算"),
            ("startability", "技術参入可能性スコア算出"),
            ("startability_ranking", "企業別/技術別ランキング"),
            ("startability_delta", "時系列変化の検出"),
            ("firm_tech_vector", "企業の技術ベクトル取得"),
            ("tech_clusters_list", "607技術クラスタの一覧"),
        ]),
        ("戦略分析 (Strategic Analysis)", [
            ("adversarial_strategy", "ゲーム理論的2社対比分析"),
            ("tech_gap", "技術ギャップ・シナジー分析"),
            ("similar_firms", "類似企業発見"),
            ("cross_domain_discovery", "異分野技術融合の発見"),
            ("invention_intelligence", "先行技術・FTO・ホワイトスペース"),
            ("sales_prospect", "ライセンス営業先の発見"),
            ("patent_market_fusion", "特許×市場シグナル統合分析"),
        ]),
        ("財務・バリュエーション (Finance & Valuation)", [
            ("patent_valuation", "特許/ポートフォリオ価値評価"),
            ("patent_option_value", "Black-Scholesリアルオプション"),
            ("tech_volatility", "技術ボラティリティ分析"),
            ("tech_beta", "CAPM技術β値"),
            ("portfolio_var", "特許ポートフォリオVaR"),
            ("bayesian_scenario", "ベイズ投資シミュレーション"),
            ("ip_due_diligence", "統合IPデューデリジェンス"),
        ]),
        ("ネットワーク分析 (Network Analysis)", [
            ("citation_network", "引用ネットワーク構築"),
            ("network_topology", "ネットワークトポロジー分析"),
            ("network_resilience", "パーコレーション理論による堅牢性"),
            ("knowledge_flow", "技術間知識フロー分析"),
            ("tech_fusion_detector", "技術融合検出"),
        ]),
        ("トレンド分析 (Trend Analysis)", [
            ("tech_trend", "技術トレンド時系列分析"),
            ("tech_trend_alert", "ホット/クーリング技術検出"),
            ("tech_entropy", "技術成熟度分析"),
            ("portfolio_evolution", "企業ポートフォリオ変遷追跡"),
        ]),
        ("M&A・投資分析 (M&A & Investment)", [
            ("ma_target", "M&A買収候補推薦"),
            ("cross_border_similarity", "国際特許類似性検出"),
        ]),
        ("SEP標準必須特許 (Standard Essential Patents)", [
            ("sep_search", "SEP宣言の検索"),
            ("sep_landscape", "標準特許ランドスケープ"),
            ("sep_portfolio", "企業のSEPポートフォリオ"),
            ("frand_analysis", "FRAND分析・ロイヤルティスタック"),
        ]),
        ("クレーム分析 (Claim Analysis)", [
            ("claim_analysis", "特許の技術スコープ分析"),
            ("claim_comparison", "複数特許の技術スコープ比較"),
            ("fto_analysis", "Freedom-to-Operate分析"),
        ]),
        ("企業グループ分析 (Corporate Group)", [
            ("corporate_hierarchy", "企業グループ構造"),
            ("group_portfolio", "グループ横断ポートフォリオ"),
            ("group_startability", "グループレベルStartability"),
        ]),
        ("AI分類 (AI Classification)", [
            ("create_category", "カスタム技術カテゴリ定義"),
            ("classify_patents", "AI特許分類"),
            ("category_landscape", "カテゴリ別ランドスケープ"),
            ("portfolio_benchmark", "カテゴリ内ベンチマーク"),
        ]),
        ("訴訟・審判分析 (Litigation & PTAB)", [
            ("ptab_search", "PTAB審判検索"),
            ("ptab_risk", "PTAB無効化リスク評価"),
            ("litigation_search", "特許訴訟検索"),
            ("litigation_risk", "訴訟リスク評価"),
        ]),
        ("モニタリング (Monitoring)", [
            ("create_watch", "監視ウォッチ作成"),
            ("list_watches", "監視ウォッチ一覧"),
            ("check_alerts", "アラート確認"),
            ("acknowledge_alerts", "アラート確認済み処理"),
            ("delete_watch", "ウォッチ削除"),
        ]),
        ("サマリー・可視化 (Summary & Visualization)", [
            ("patent_summary", "特許サマリー生成"),
            ("technology_brief", "技術ブリーフ作成"),
            ("tech_map", "技術マップ(Mermaid)"),
            ("citation_graph_viz", "引用グラフ可視化"),
            ("firm_landscape", "企業ランドスケープ"),
            ("startability_heatmap", "Startabilityヒートマップ"),
        ]),
    ]

    for cat_name, tools in categories:
        add_heading_styled(doc, cat_name, level=2)
        add_styled_table(doc,
            ["ツール名", "説明"],
            [[t[0], t[1]] for t in tools],
        )
        doc.add_paragraph()

    add_page_break(doc)


def build_example(doc, num, title, use_case, tool_call, params, result_text,
                  table_data=None, insight=""):
    """Build a single execution example."""
    add_heading_styled(doc, f"Example {num}: {title}", level=2)

    add_body_text(doc, f"ユースケース: {use_case}")

    # Tool call
    styled_paragraph(doc, "ツール呼び出し:", font_size=Pt(11), bold=True,
                     color=DARK_BLUE, space_after=Pt(4))
    add_code_block(doc, f"{tool_call}({params})")

    # Result
    styled_paragraph(doc, "実行結果:", font_size=Pt(11), bold=True,
                     color=DARK_BLUE, space_after=Pt(4))

    if table_data:
        headers, rows = table_data
        add_styled_table(doc, headers, rows)
        doc.add_paragraph()

    if result_text:
        for line in result_text.split("\n"):
            if line.strip():
                add_bullet(doc, line.strip())

    if insight:
        doc.add_paragraph()
        styled_paragraph(doc, "ビジネスインサイト:", font_size=Pt(11), bold=True,
                         color=DARK_BLUE, space_after=Pt(4))
        add_body_text(doc, insight)


def build_examples(doc):
    """Pages 10-25: 10 Real Execution Examples."""
    add_heading_styled(doc, "実行例: 10のリアルデータ分析", level=1)

    add_body_text(
        doc,
        "以下は、Patent Space MCPの実データを用いた分析例です。"
        "すべて実際のツール実行結果に基づいています。"
    )

    # Example 1
    build_example(
        doc, 1,
        "トヨタ自動車の技術ポートフォリオ",
        "トヨタ自動車の特許ポートフォリオの全体像を把握する",
        "firm_tech_vector",
        'firm_query="トヨタ自動車"',
        "patent_count: 260,223件\n"
        "dominant_cpc: B60W（車両制御システム）\n"
        "tech_diversity: 3.91（Shannon Entropy）\n"
        "tech_concentration: 0.29\n"
        "技術ベクトル: 64次元ベクトルとして取得可能",
        None,
        "トヨタは26万件超の特許を保有し、車両制御(B60W)が最も集中する技術領域です。"
        "tech_diversity 3.91は広範な技術分散を示しており、"
        "自動車製造に留まらない幅広い技術投資戦略を反映しています。"
        "tech_concentration 0.29は、特定技術への過度な集中がないことを示します。"
    )

    add_page_break(doc)

    # Example 2
    build_example(
        doc, 2,
        "ソニーのAI技術参入可能性",
        "ソニーグループがAI/機械学習領域にどの程度参入可能かを定量評価",
        "startability",
        'firm_query="ソニーグループ", tech_query_or_cluster_id="G06N_0"',
        "Startability Score: 0.935（Very High）\n"
        "gate_open: true（参入可能）\n"
        "phi_tech_cosine: 0.771\n"
        "phi_tech_distance: 0.680",
        None,
        "ソニーグループはAI/機械学習(G06N)領域に対して極めて高い参入可能性(0.935)を持ちます。"
        "既存の画像処理・センサー技術の特許ポートフォリオがAI技術と高い親和性を持つことが、"
        "phi_tech_cosine 0.771に反映されています。gate_open=trueは、"
        "この技術領域への新規参入が現実的に可能であることを示しています。"
    )

    doc.add_paragraph()

    # Example 3
    build_example(
        doc, 3,
        "電池技術(H01M)の出願トレンド",
        "全固体電池を含む電池技術の出願動向を時系列で分析",
        "tech_trend",
        'query="全固体電池", cpc_prefix="H01M"',
        None,
        (["年", "出願件数", "前年比"],
         [
             ["2016", "6,941", "-"],
             ["2017", "7,644", "+10.1%"],
             ["2018", "8,860", "+15.9%"],
             ["2019", "8,706", "-1.7%"],
             ["2020", "9,756", "+12.1%"],
             ["2021", "10,786", "+10.6% (PEAK)"],
             ["2022", "10,626", "-1.5%"],
             ["2023", "8,395", "-21.0%"],
             ["2024", "2,903", "-65.4% (蓄積中)"],
         ]),
        "電池技術(H01M)の出願は2021年にピーク(10,786件)を迎え、その後減少傾向にあります。"
        "2024年のデータは蓄積途中ですが、2023年の21%減少は注目すべきシグナルです。"
        "全固体電池の基盤技術が成熟期に入りつつある可能性を示唆しています。"
    )

    add_page_break(doc)

    # Example 4
    build_example(
        doc, 4,
        "FTO分析（固体電池硫化物電解質）",
        "硫化物系固体電解質を用いた全固体電池のFreedom-to-Operate評価",
        "fto_analysis",
        'text="solid state battery using sulfide electrolyte", cpc_codes=["H01M10"]',
        "overall_risk: LOW\n"
        "risk_score: 0\n"
        "Recommendation: FTOリスクは比較的低い領域です。ただし、新規出願の動向には引き続き注意が必要です。",
        None,
        "硫化物系固体電解質のFTOリスクは低いと評価されました。"
        "これは当該技術領域のブロッキング特許が少ない、または期限切れ・失効している可能性を示唆します。"
        "ただし、本分析は予備的なものであり、実際のFTO意見書は専門の弁理士による精査が必要です。"
        "\n\n※ Disclaimer: Professional patent attorney review required for actual FTO opinions."
    )

    doc.add_paragraph()

    # Example 5
    build_example(
        doc, 5,
        "電池技術の特許価値評価",
        "電池技術(H01M)領域全体の特許価値を定量評価",
        "patent_valuation",
        'query="H01M", query_type="technology"',
        None,
        (["指標", "値"],
         [
             ["平均特許RFR-NPV", "167.5万円"],
             ["領域合計価値", "112,090万円"],
             ["ロイヤルティレート（業界標準）", "2-6%（典型 4%）"],
             ["オプション価値", "124.5万円"],
             ["原資産価値 (S)", "201.1万円"],
             ["行使価格 (K)", "140.7万円"],
             ["ボラティリティ (σ)", "0.378"],
             ["成長率", "-10.9%"],
         ]),
        "電池技術の平均特許価値は167.5万円/件で、領域全体では約11.2億円規模です。"
        "Black-Scholesモデルによるオプション価値124.5万円は、将来の商業化オプションとしての"
        "潜在価値を反映しています。ボラティリティ0.378は技術変動がある程度高いことを示し、"
        "成長率-10.9%は技術の成熟化傾向を裏付けています。"
    )

    add_page_break(doc)

    # Example 6
    build_example(
        doc, 6,
        "ソニーのIPデューデリジェンス",
        "ソニーへのグロース投資検討のためのIP精査",
        "ip_due_diligence",
        'target_firm="sony", investment_type="growth"',
        None,
        (["指標", "値"],
         [
             ["総特許数", "142,768件"],
             ["Fusion Score", "15.3 / 100"],
             ["技術多様性", "3.24 (entropy)"],
             ["主要技術領域", "H01L（半導体デバイス）"],
         ]),
        "ソニーは14.2万件の特許を保有する大規模ポートフォリオを持ちます。"
        "主要技術はH01L(半導体デバイス)で、イメージセンサー技術の強さが反映されています。"
        "技術多様性3.24は幅広い技術領域をカバーしていることを示し、"
        "グロース投資の観点から安定的な技術基盤があると評価できます。"
    )

    doc.add_paragraph()

    # Example 7
    build_example(
        doc, 7,
        "トヨタのStartabilityランキング",
        "トヨタが参入可能な技術領域をスコア順にランキング",
        "startability_ranking",
        'mode="by_firm", query="トヨタ自動車"',
        None,
        (["順位", "技術クラスタ", "Startability"],
         [
             ["1", "Y02T_0 (気候変動緩和・輸送)", "0.986"],
             ["2", "B60K_0 (車両推進装置)", "0.983"],
             ["3", "B60Y_0 (車両分類)", "0.983"],
             ["4", "B60L_0 (電気車両推進)", "0.980"],
             ["5", "B60W_0 (車両制御)", "0.975"],
             ["6", "F02N_0 (エンジン始動)", "0.968"],
             ["7", "B62D_0 (車両シャーシ)", "0.967"],
             ["8", "B60R_0 (車両装備品)", "0.966"],
             ["9", "Y10T_0 (技術主題)", "0.961"],
             ["10", "G05D_0 (制御システム)", "0.960"],
         ]),
        "トヨタの最高Startabilityは気候変動緩和・輸送(Y02T_0)で0.986。"
        "上位10領域中8つが車両関連(B60x)であり、コア技術領域での圧倒的な参入優位性を示します。"
        "注目すべきは10位のG05D_0(制御システム)で、自動運転やロボティクスへの展開基盤を持つことを示唆します。"
    )

    add_page_break(doc)

    # Example 8
    build_example(
        doc, 8,
        "5G SEP標準特許ランドスケープ",
        "5G NR標準の特許宣言状況を分析",
        "sep_landscape",
        'standard="5G"',
        "総宣言数: 950件\n"
        "宣言企業数: 19社",
        (["順位", "宣言企業", "宣言数"],
         [
             ["1", "Qualcomm", "50"],
             ["2", "Samsung", "50"],
             ["3", "Sony", "50"],
             ["4", "Panasonic", "50"],
             ["5", "LG", "50"],
         ]),
        "5G SEP宣言は950件、19社が参加しています。"
        "Qualcomm、Samsung、Sony、Panasonic、LGが上位を占め、"
        "FRAND条件でのライセンス交渉において、これらの企業が大きな交渉力を持つことがわかります。"
    )

    doc.add_paragraph()

    # Example 9
    build_example(
        doc, 9,
        "SamsungのSEPポートフォリオ",
        "Samsungが宣言する標準必須特許の全体像",
        "sep_portfolio",
        'firm_query="Samsung"',
        None,
        (["指標", "値"],
         [
             ["総宣言数", "450件"],
             ["カバー標準数", "9標準"],
             ["対象標準", "5G NR, AVC, Bluetooth, HEVC, LTE, NFC, UMTS, Wi-Fi, Wi-Fi 6"],
             ["ピーク年", "2018年（45宣言）"],
         ]),
        "Samsungは9つの通信標準にわたり450件のSEP宣言を行っており、"
        "通信分野での技術リーダーシップを確立しています。"
        "5G NRからBluetoothまで幅広い標準をカバーしており、"
        "IoTからモバイル通信まで多層的なライセンスポジションを構築しています。"
    )

    add_page_break(doc)

    # Example 10
    build_example(
        doc, 10,
        "特許のクレーム分析",
        "個別特許の技術スコープを詳細分析",
        "claim_analysis",
        'publication_number="JP-7637366-B1"',
        None,
        (["項目", "値"],
         [
             ["特許番号", "JP-7637366-B1"],
             ["タイトル", "流体用フィルタ"],
             ["技術要素", "側壁構造、分離壁、乱流発生機構、フィルタ領域"],
             ["スコープ", "narrow（狭い）"],
             ["出願人", "大学共同利用機関法人自然科学研究機構, 学校法人君が淵学園"],
             ["出願日", "2024-09-13"],
         ]),
        "JP-7637366-B1は流体用フィルタに関する特許で、大学・学校法人による共同出願です。"
        "技術スコープは「narrow」と評価され、側壁構造や乱流発生機構など具体的な構造要素に"
        "限定されたクレーム範囲を持ちます。学術機関由来の特許であり、"
        "ライセンスアウトの可能性を検討する際の基礎情報として有用です。"
    )

    add_page_break(doc)


def build_architecture(doc):
    """Pages 26-28: Architecture."""
    add_heading_styled(doc, "アーキテクチャ", level=1)

    add_body_text(
        doc,
        "Patent Space MCPは、高性能なデータベースバックエンドと"
        "MCPプロトコルサーバーを組み合わせたアーキテクチャで構成されています。"
    )

    add_heading_styled(doc, "システム構成図", level=2)

    # Architecture as Word table (not ASCII art which breaks in Word)
    arch_table = doc.add_table(rows=5, cols=4)
    arch_table.style = "Table Grid"
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: AI Client layer
    cell = arch_table.cell(0, 0)
    cell.merge(arch_table.cell(0, 3))
    cell.text = "AI Assistant (Claude Desktop / Claude Code / API Client)"
    set_cell_shading(cell, HEADER_BG)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(11)

    # Row 1: Protocol
    cell = arch_table.cell(1, 0)
    cell.merge(arch_table.cell(1, 3))
    cell.text = "MCP Protocol (Streamable HTTP) — https://patent-space.dev/mcp"
    set_cell_shading(cell, "E8F0FE")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(10)

    # Row 2: MCP Server
    cell = arch_table.cell(2, 0)
    cell.merge(arch_table.cell(2, 3))
    cell.text = "MCP Server — FastMCP 2.14.5 (Python 3.13) — 67 Tools"
    set_cell_shading(cell, HEADER_BG)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(11)

    # Row 3: Data layers
    data_layers = [
        ("patents\n150M+ records", HEADER_BG),
        ("firm_tech_vectors\n27.8K entries", ALT_ROW_BG),
        ("startability_surface\n10.2M rows", HEADER_BG),
        ("SEP/PTAB/Litigation\n83K+ records", ALT_ROW_BG),
    ]
    for i, (txt, bg) in enumerate(data_layers):
        cell = arch_table.cell(3, i)
        cell.text = txt
        set_cell_shading(cell, bg)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)

    # Row 4: Database
    cell = arch_table.cell(4, 0)
    cell.merge(arch_table.cell(4, 3))
    cell.text = "SQLite Database — 650GB+ (HDD RAID1, NVMe cache)"
    set_cell_shading(cell, "1B365D")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(11)

    add_heading_styled(doc, "サーバー仕様", level=2)

    add_styled_table(doc,
        ["項目", "仕様"],
        [
            ["フレームワーク", "FastMCP 2.14.5 (Python 3.13)"],
            ["プロトコル", "Model Context Protocol (MCP)"],
            ["トランスポート", "Streamable HTTP / SSE"],
            ["エンドポイント", "https://patent-space.dev/mcp"],
            ["コンテナ", "Docker (48GB memory limit)"],
            ["データベース", "SQLite 650GB+ (150M+ patents)"],
            ["ヘルスチェック", "https://patent-space.dev/health"],
            ["MCPエンドポイント", "https://patent-space.dev/mcp"],
        ]
    )

    doc.add_paragraph()

    add_heading_styled(doc, "パフォーマンス最適化", level=2)

    optimizations = [
        "事前計算テーブル: startability_surface (10.2M行)、firm_tech_vectors (27.8K行) による即時応答",
        "コネクションプーリング: threading.local()によるスレッドごとの永続接続",
        "FTS5全文検索: 高速テキスト検索エンジン内蔵",
        "バッチクエリ: N+1問題を排除し、IN句によるバッチ取得を実施",
        "ページキャッシュ事前ウォーム: バックグラウンドスレッドによるホットテーブルの事前読込",
        "高速パスランキング: 事前計算テーブルのみを使用する _fast_batch_rank() 実装",
    ]
    for opt in optimizations:
        add_bullet(doc, opt)

    add_page_break(doc)


def build_connection_guide(doc):
    """Pages 29-30: Connection Guide."""
    add_heading_styled(doc, "接続方法", level=1)

    add_heading_styled(doc, "Claude Desktop からの接続", level=2)

    add_body_text(doc, "Claude Desktopの設定画面から簡単に接続できます:")

    # Step-by-step as a table
    steps_data = [
        ["1", "Claude Desktopを開く"],
        ["2", "設定 (Settings) → 開発者 (Developer) を開く"],
        ["3", "「カスタムコネクタ」(Custom Connector) をクリック"],
        ["4", "URL欄に https://patent-space.dev/mcp を入力"],
        ["5", "「接続」をクリックして完了"],
    ]
    add_styled_table(doc, ["手順", "操作"], steps_data)

    doc.add_paragraph()
    add_body_text(doc, "または、~/.claude/mcp.json に以下の設定を追加することもできます:")

    add_code_block(doc,
        '{\n'
        '  "mcpServers": {\n'
        '    "patent-space": {\n'
        '      "type": "streamable-http",\n'
        '      "url": "https://patent-space.dev/mcp"\n'
        '    }\n'
        '  }\n'
        '}'
    )

    add_heading_styled(doc, "Direct HTTP (curl)", level=2)

    add_code_block(doc,
        '# ヘルスチェック\n'
        'curl https://patent-space.dev/health\n'
        '\n'
        '# ツール呼び出し例\n'
        'curl -X POST https://patent-space.dev/mcp \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{\n'
        '    "jsonrpc": "2.0",\n'
        '    "id": 1,\n'
        '    "method": "tools/call",\n'
        '    "params": {\n'
        '      "name": "patent_search",\n'
        '      "arguments": {\n'
        '        "query": "全固体電池",\n'
        '        "max_results": 5\n'
        '      }\n'
        '    }\n'
        '  }\''
    )

    add_heading_styled(doc, "Python SDK", level=2)

    add_code_block(doc,
        'from mcp import ClientSession\n'
        'from mcp.client.streamable_http import streamablehttp_client\n'
        '\n'
        'async def main():\n'
        '    async with streamablehttp_client("https://patent-space.dev/mcp") as (\n'
        '        read, write, _\n'
        '    ):\n'
        '        async with ClientSession(read, write) as session:\n'
        '            await session.initialize()\n'
        '            result = await session.call_tool(\n'
        '                "patent_search",\n'
        '                arguments={"query": "AI", "max_results": 10}\n'
        '            )\n'
        '            print(result)'
    )

    add_page_break(doc)


def build_roadmap(doc):
    """Pages 31-33: Roadmap."""
    add_heading_styled(doc, "ロードマップ", level=1)

    # Phase 1
    add_heading_styled(doc, "Phase 1: 日本特許基盤（完了）", level=2)
    phase1 = [
        "67の専門分析ツール実装・稼働",
        "グローバル特許 150,000,000件超のメタデータ格納",
        "607技術クラスタの定義・マッピング",
        "10,267,405件のStartabilityスコア事前計算",
        "20,736件の企業名寄せエントリ構築",
        "SEP宣言データ (950+件) 統合",
        "PTAB審判 (7,605件) / 訴訟 (74,629件) データ統合",
        "GDELTメディアシグナル (46企業) 統合",
    ]
    for item in phase1:
        add_bullet(doc, item)

    doc.add_paragraph()

    # Phase 2
    add_heading_styled(doc, "Phase 2: グローバル特許拡張（92%完了）", level=2)
    phase2 = [
        "グローバル150M件超特許の取込（4,640/5,000ファイル完了）",
        "USPTO / EPO / WIPO / CN / KR 特許の統合",
        "クロスボーダー分析の本格化",
        "多言語対応の強化（英語、中国語、韓国語）",
        "企業名寄せの国際拡張（S&P 500、FTSE 100等）",
    ]
    for item in phase2:
        add_bullet(doc, item)

    doc.add_paragraph()

    # Phase 3
    add_heading_styled(doc, "Phase 3: 引用ネットワーク最適化（計画中）", level=2)
    phase3 = [
        "大規模引用グラフの最適化（BFS深度2+）",
        "リアルタイムアラート機能の実装",
        "カスタムダッシュボード機能",
        "API rate limitingとマルチテナント対応",
        "WebUIの構築",
    ]
    for item in phase3:
        add_bullet(doc, item)

    doc.add_paragraph()

    # Phase 4
    add_heading_styled(doc, "Phase 4: 次世代機能（将来構想）", level=2)
    phase4 = [
        "多言語クレーム分析（日英中韓）",
        "クロスジュリスディクションFTO自動分析",
        "特許出願ドラフト支援",
        "リアルタイム特許出願監視",
        "LLMによる特許品質自動評価",
        "ブロックチェーンベースの特許ライセンス管理",
    ]
    for item in phase4:
        add_bullet(doc, item)

    add_page_break(doc)


def build_comparison(doc):
    """Pages 34-35: Competitive Comparison."""
    add_heading_styled(doc, "競合比較", level=1)

    add_body_text(
        doc,
        "Patent Space MCPと主要商用特許分析プラットフォームの機能比較です。"
    )

    headers = ["機能", "Patent Space MCP", "PatSnap (300万円/年)", "Derwent (200万円/年)"]
    rows = [
        ["ツール数", "67", "約30", "約20"],
        ["日本特許", "26.2M件", "対応", "対応"],
        ["グローバル特許", "150M件超", "170M+件", "100M+件"],
        ["Startability分析", "対応 (独自)", "非対応", "非対応"],
        ["MCP対応", "対応", "非対応", "非対応"],
        ["AI分類", "対応 (カスタム)", "対応", "一部対応"],
        ["ゲーム理論分析", "対応", "非対応", "非対応"],
        ["リアルオプション", "対応", "非対応", "非対応"],
        ["PTAB/訴訟分析", "対応", "対応", "一部対応"],
        ["SEP分析", "対応", "一部対応", "対応"],
        ["ベイズ投資分析", "対応", "非対応", "非対応"],
        ["技術β/VaR", "対応", "非対応", "非対応"],
        ["企業グループ分析", "対応", "一部対応", "一部対応"],
        ["モニタリング", "対応", "対応", "対応"],
        ["価格", "オープンソース", "300万円/年", "200万円/年"],
    ]

    table = add_styled_table(doc, headers, rows)

    doc.add_paragraph()

    add_heading_styled(doc, "独自機能ハイライト", level=2)

    add_body_text(
        doc,
        "Patent Space MCPには、商用DBにはない独自の分析機能が搭載されています:"
    )

    unique_features = [
        "Startability分析: 企業×技術の参入可能性を定量スコア化（10.2M件の事前計算）。"
        "企業が「次にどの技術領域に参入できるか」を科学的に評価する世界初の機能。",
        "ゲーム理論的戦略分析: 2社間の特許ポートフォリオを攻防の観点から分析。"
        "攻撃対象、防御優先技術、先制機会を特定。",
        "Black-Scholesリアルオプション: 特許をリアルオプションとして評価。"
        "デルタ、シータ、ベガ等のGreeksを含む本格的なオプション分析。",
        "技術β/VaR: CAPM理論を特許に適用。技術のマーケット感応度とリスク量を定量化。",
        "ベイズ投資シミュレーション: データ駆動の事前分布を構築し、"
        "ユーザーの私的情報をベイズ更新で統合するモンテカルロシミュレーション。",
        "MCPネイティブ: AI Assistant (Claude, Cursor等)から直接67ツールを呼び出し可能。"
        "自然言語で「トヨタのAI技術への参入可能性を分析して」と指示するだけで完了。",
    ]
    for feat in unique_features:
        add_bullet(doc, feat)

    add_page_break(doc)


def setup_document_styles(doc):
    """Configure document-wide styles, margins, and page numbers."""
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_BODY)

    # Customize heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f"Heading {i}"]
        heading_style.font.name = FONT_BODY
        heading_style.font.color.rgb = DARK_BLUE
        hPr = heading_style.element.get_or_add_rPr()
        hFonts = hPr.find(qn("w:rFonts"))
        if hFonts is None:
            hFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            hPr.insert(0, hFonts)
        hFonts.set(qn("w:eastAsia"), FONT_BODY)

    # Add page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Page number field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

        # Add header with "Confidential"
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Confidential")
        hrun.font.name = FONT_BODY
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        hrun.font.italic = True


def build_toc_placeholder(doc):
    """Add a Table of Contents placeholder."""
    add_heading_styled(doc, "目次", level=1)

    add_body_text(
        doc,
        "※ Wordで開いた後、目次フィールドを右クリック → 「フィールドの更新」で"
        "ページ番号付き目次が自動生成されます。"
    )

    # Insert TOC field
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u '
        '</w:instrText>'
    )
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar2)

    run4 = p.add_run("[ 目次を更新するには、ここを右クリック → フィールドの更新 ]")
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar3)

    add_page_break(doc)


def build_appendix(doc):
    """Final page: Appendix / Contact."""
    add_heading_styled(doc, "付録", level=1)

    add_heading_styled(doc, "用語集", level=2)

    add_styled_table(doc,
        ["用語", "説明"],
        [
            ["MCP", "Model Context Protocol - AIモデルがツールを呼び出すための標準プロトコル"],
            ["CPC", "Cooperative Patent Classification - 国際特許分類体系"],
            ["Startability", "企業が特定技術領域に参入する可能性・準備度を定量化した独自指標"],
            ["SEP", "Standard Essential Patent - 技術標準の実装に不可欠な特許"],
            ["FRAND", "Fair, Reasonable and Non-Discriminatory - SEPのライセンス条件"],
            ["FTO", "Freedom to Operate - 特許侵害リスクなく事業活動できるかの分析"],
            ["PTAB", "Patent Trial and Appeal Board - 米国特許審判部"],
            ["VaR", "Value at Risk - 特許ポートフォリオのリスク量"],
            ["phi_tech", "技術適合度 - 企業の技術ベクトルと対象技術クラスタの類似度"],
            ["Shannon Entropy", "技術多様性の指標。値が大きいほど多様な技術に分散"],
            ["GDELT", "Global Database of Events, Language, and Tone - メディアシグナルDB"],
            ["FTS5", "Full-Text Search 5 - SQLite組み込みの全文検索エンジン"],
        ]
    )

    doc.add_paragraph()

    add_heading_styled(doc, "技術スタック", level=2)

    add_styled_table(doc,
        ["コンポーネント", "技術"],
        [
            ["MCPフレームワーク", "FastMCP 2.14.5"],
            ["言語", "Python 3.13"],
            ["データベース", "SQLite (310GB+)"],
            ["コンテナ", "Docker"],
            ["サーバー", "Hetzner Dedicated (Intel i7-8700, 64GB RAM)"],
            ["データソース", "Google Patents Public Data (BigQuery)"],
            ["テキスト検索", "FTS5"],
            ["プロトコル", "MCP (Streamable HTTP / SSE)"],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph()

    styled_paragraph(
        doc,
        "Patent Space MCP - Confidential Product Guide v1.0",
        font_size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    styled_paragraph(
        doc,
        "Generated: 2026-03-09",
        font_size=Pt(9), color=RGBColor(0x99, 0x99, 0x99),
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )


# ── Main ───────────────────────────────────────────────────────────

def main():
    """Generate the complete Patent Space MCP Product Guide."""
    print("Generating Patent Space MCP Product Guide...")

    doc = Document()

    # Configure styles and page setup
    setup_document_styles(doc)

    # Build all sections
    build_title_page(doc)
    build_toc_placeholder(doc)
    build_executive_summary(doc)
    build_data_foundation(doc)
    build_tool_list(doc)
    build_examples(doc)
    build_architecture(doc)
    build_connection_guide(doc)
    build_roadmap(doc)
    build_comparison(doc)
    build_appendix(doc)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    print("Done.")


if __name__ == "__main__":
    main()
